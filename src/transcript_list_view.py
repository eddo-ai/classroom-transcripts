import streamlit as st
from utils.view_table import get_table_client, list_table_items
from datetime import datetime
import pytz
import assemblyai as aai
import os
import logging
from docx import Document
from io import BytesIO
import pydantic

from utils.azure_storage import get_sas_url_for_audio_file_name

# Initialize logging with more detailed format
logger = logging.getLogger(__name__)

# Wrap initialization in try-except
try:
    DEBUG = bool(st.secrets.get("DEBUG", False))
    table_name = st.session_state.get(
        "table_name", st.secrets.get("AZURE_STORAGE_TABLE_NAME")
    )
    TRANSCRIPT_PREVIEW_MAX_LENGTH = 1000
    TRANSCRIPT_PREVIEW_SPEAKER_TURNS = 5

    # Initialize AssemblyAI client
    api_key = os.getenv("ASSEMBLYAI_API_KEY")
    if not api_key:
        logger.error("AssemblyAI API key not found in environment")
        st.error("Configuration error. Please contact support.")
        st.stop()
    aai.settings.api_key = api_key
    transcriber = aai.Transcriber()

    if not st.experimental_user.get("is_logged_in"):
        st.info("Please log in to continue")
        st.login()
        st.stop()

    # Initialize session state for status values if not already set
    if "transcription_statuses" not in st.session_state:
        st.session_state["transcription_statuses"] = [
            "queued",
            "processing",
            "completed",
            "error",
            "failed",
        ]

    # US timezone options
    US_TIMEZONES = [
        "US/Eastern",
        "US/Central",
        "US/Mountain",
        "US/Pacific",
        "US/Alaska",
        "US/Hawaii",
    ]

    # Initialize timezone in session state
    if "timezone" not in st.session_state:
        st.session_state["timezone"] = "US/Pacific"

    # Initialize table client early
    try:
        table_client = get_table_client(table_name)
        if not table_client:
            raise ValueError("Could not initialize table client")
    except Exception as e:
        logger.error(f"Table client initialization failed: {str(e)}", exc_info=True)
        st.error("Database connection error. Please try again later.")
        st.stop()

    # Initialize session state for pagination
    if "items_per_page" not in st.session_state:
        st.session_state["items_per_page"] = 5  # Initial number of items to show
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = 1

    # Get admin emails from Streamlit secrets
    ADMIN_EMAILS = [
        email.strip().lower() for email in st.secrets.get("admin_emails", "").split(",")
    ]

    # Debug logging for admin list
    if DEBUG:
        st.write("Debug - Admin emails:", ADMIN_EMAILS)

    def is_admin(email: str) -> bool:
        """Check if the given email belongs to an admin"""
        return email.lower() in ADMIN_EMAILS

    def reset_pagination():
        """Reset pagination state"""
        st.session_state["items_per_page"] = 5
        st.session_state["current_page"] = 1

    def format_file_size(size_in_bytes):
        """Convert bytes to human readable format"""
        if not isinstance(size_in_bytes, (int, float)):
            return "N/A"
        for unit in ["B", "KB", "MB", "GB"]:
            if size_in_bytes < 1024:
                return f"{size_in_bytes:.2f} {unit}"
            size_in_bytes /= 1024
        return f"{size_in_bytes:.2f} TB"

    def localized_timestamp(timestamp):
        """Get localized timestamp"""
        local_timestamp = timestamp.astimezone(
            pytz.timezone(st.session_state["timezone"])
        )
        time_string = local_timestamp.strftime("%Y-%m-%d %H:%M:%S %Z")
        return time_string

    def get_transcript_status(transcript_id):
        """Get transcript status from AssemblyAI"""
        # Skip test data
        if transcript_id.startswith("test_"):
            return "completed"

        try:
            transcript = aai.Transcript.get_by_id(transcript_id)
            return transcript.status.value
        except Exception as e:
            if "not found" in str(e).lower():
                return "error"  # Transcript doesn't exist in AssemblyAI
            st.error(f"Error getting transcript status: {str(e)}")
            return "error"

    def generate_transcript_markdown(
        transcript, max_length=None, max_speaker_turns=None
    ):
        """
        Generate markdown formatted text from an AssemblyAI transcript.

        Args:
            transcript: AssemblyAI Transcript object
            max_length (int, optional): Maximum length of text to include
            max_speaker_turns (int, optional): Maximum number of speaker turns to include

        Returns:
            str: Markdown formatted transcript text
        """
        logger.debug(
            "Starting transcript markdown generation",
            extra={
                "has_transcript": bool(transcript),
                "max_length": max_length,
                "max_speaker_turns": max_speaker_turns,
            },
        )

        if not transcript:
            logger.warning("No transcript provided to generate_transcript_markdown")
            return "No transcript available"

        markdown_lines = []

        try:
            # Handle transcripts with speaker detection
            if hasattr(transcript, "utterances") and transcript.utterances:
                logger.debug(
                    "Processing transcript with speaker detection",
                    extra={"utterance_count": len(transcript.utterances)},
                )

                for i, utterance in enumerate(transcript.utterances):
                    # Break if we've hit the max speaker turns
                    if max_speaker_turns and i >= max_speaker_turns:
                        markdown_lines.append(
                            "\n*[Additional transcript content truncated...]*"
                        )
                        logger.debug(
                            f"Reached max speaker turns limit: {max_speaker_turns}"
                        )
                        break

                    try:
                        # Format timestamp as [00:00:00]
                        start_seconds = (
                            utterance.start / 1000.0
                        )  # Convert milliseconds to seconds
                        hours = int(start_seconds // 3600)
                        minutes = int((start_seconds % 3600) // 60)
                        seconds = int(start_seconds % 60)
                        timestamp = f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

                        # Format as [timestamp] **Speaker X**: text
                        speaker_letter = (
                            chr(65 + (utterance.speaker - 1))
                            if isinstance(utterance.speaker, int)
                            else utterance.speaker
                        )
                        speaker_text = f"{timestamp} **Speaker {speaker_letter}**: {utterance.text}"
                        markdown_lines.append(speaker_text)

                        # Check total length if max_length specified
                        current_text = "\n".join(markdown_lines)
                        if max_length and len(current_text) >= max_length:
                            truncate_length = max_length - len(
                                "\n\n*[Additional transcript content truncated...]*"
                            )
                            markdown_lines[-1] = str(markdown_lines[-1])[
                                :truncate_length
                            ]
                            markdown_lines.append(
                                "\n*[Additional transcript content truncated...]*"
                            )
                            logger.debug(f"Reached max length limit: {max_length}")
                            break
                    except Exception as e:
                        logger.error(
                            f"Error processing utterance {i}",
                            extra={
                                "error": str(e),
                                "utterance": str(utterance),
                            },
                            exc_info=True,
                        )
                        continue

            # Handle transcripts without speaker detection
            elif hasattr(transcript, "text") and transcript.text:
                logger.debug("Processing transcript without speaker detection")
                text = transcript.text
                if max_length:
                    truncate_length = max_length
                    text = str(text)[:truncate_length] + (
                        "..." if len(transcript.text) > truncate_length else ""
                    )
                    logger.debug(f"Truncated text to length: {truncate_length}")
                markdown_lines.append(text)
            else:
                logger.warning(
                    "Transcript has neither utterances nor text",
                    extra={"transcript_attrs": dir(transcript)},
                )
                return "Transcript format not recognized"

            result = "\n\n".join(markdown_lines)
            logger.debug(
                "Completed transcript markdown generation",
                extra={"result_length": len(result)},
            )
            return result

        except Exception as e:
            logger.error(
                "Failed to generate transcript markdown",
                extra={"error": str(e)},
                exc_info=True,
            )
            return f"Error generating transcript: {str(e)}"

    def generate_transcript_docx(transcript):
        """
        Generate a docx file from an AssemblyAI transcript.

        Args:
            transcript: AssemblyAI Transcript object

        Returns:
            bytes: The generated docx file as bytes
        """
        doc = Document()
        doc.add_heading("Transcript", 0)

        if transcript.utterances:
            # Add each speaker's text as a paragraph
            for utterance in transcript.utterances:
                p = doc.add_paragraph()
                # Add speaker label in bold
                speaker_run = p.add_run(f"Speaker {utterance.speaker}: ")
                speaker_run.bold = True
                # Add the text
                p.add_run(utterance.text)
                # Add spacing between utterances
                p.add_run("\n")
        else:
            # Add the full text as a single paragraph
            doc.add_paragraph(transcript.text)

        # Add metadata
        doc.core_properties.title = "Transcript"
        doc.core_properties.comments = "Generated from AssemblyAI transcription"

        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()

    st.title("🔍 Audio Files & Transcriptions")

    if DEBUG:
        st.write(
            f"Debug - Table name: {st.session_state.get('table_name', 'Transcriptions')}"
        )

    def can_view_transcript(transcript_email: str, user_email: str) -> bool:
        """Check if user can view a specific transcript"""
        if is_admin(user_email):
            return True
        # If no uploader email is set, only admins can view
        if not transcript_email:
            return False
        return transcript_email.lower() == user_email.lower()

    def get_transcript_statuses():
        """Get all transcript statuses in one API call with timeout protection"""
        try:
            logger.info("Starting batch transcript status fetch")
            status_map = {}
            params = aai.ListTranscriptParameters(limit=100)

            try:
                page = transcriber.list_transcripts(params)
                for t in page.transcripts:
                    if t.id.startswith("test_"):
                        status_map[t.id] = "completed"
                    else:
                        status_map[t.id] = t.status.value
                logger.info(
                    "Successfully fetched transcript statuses",
                    extra={
                        "count": len(status_map),
                        "statuses": {
                            k: v
                            for k, v in status_map.items()
                            if not k.startswith("test_")
                        },
                    },
                )
            except Exception as e:
                if "api error" in str(e).lower():
                    logger.error(
                        "AssemblyAI API error during status fetch",
                        exc_info=True,
                        extra={"error": str(e)},
                    )
                else:
                    logger.error(
                        "Error fetching transcript statuses",
                        exc_info=True,
                        extra={"error": str(e)},
                    )
                return {}

            return status_map
        except Exception as e:
            logger.error(
                "Unexpected error in get_transcript_statuses",
                exc_info=True,
                extra={"error": str(e)},
            )
            return {}

    def query_table_entities(table_client, user_email: str):
        """
        Query table entities based on user permissions.

        Args:
            table_client: Azure TableClient instance
            user_email: Email of the current user

        Returns:
            List of entities the user has permission to view
        """
        if not user_email:
            logger.warning("Attempted query with empty user email")
            return []

        try:
            logger.info(
                "Starting table query",
                extra={"user_email": user_email, "is_admin": is_admin(user_email)},
            )

            if not is_admin(user_email):
                logger.debug(
                    "Fetching items for user", extra={"user_email": user_email}
                )
                filter_condition = f"uploaderEmail eq '{user_email.lower()}'"
                items = list(table_client.query_entities(filter_condition))
            else:
                logger.debug("Fetching all items (admin user)")
                items = list_table_items(
                    st.session_state.get(
                        "table_name", st.secrets.get("AZURE_STORAGE_TABLE_NAME")
                    )
                )

            logger.info(
                "Table query completed",
                extra={
                    "items_count": len(items) if items else 0,
                    "user_email": user_email,
                },
            )
            return items

        except Exception as e:
            logger.error(
                "Table query error",
                exc_info=True,
                extra={"error": str(e), "user_email": user_email},
            )
            return []

    def load_table_data(_table_client):
        """Load and process table data with caching"""
        MIN_DATE = datetime(2000, 1, 1, tzinfo=pytz.UTC)

        logger.info("Starting table data load")

        user = st.experimental_user
        validated_email = user.get("email") if user.get("email_verified") else None

        if validated_email is None:
            logger.warning("No validated email available for user")
            return []

        # Use consolidated query function
        items = query_table_entities(_table_client, str(validated_email))

        if not items:
            logger.info("No items found in table query")
            return []

        logger.info("Fetching transcript statuses", extra={"items_count": len(items)})

        # Get all transcript statuses at once
        transcript_statuses = get_transcript_statuses()
        items_list = []

        for item in items:
            try:
                item_dict = dict(item)

                # Add formatted size
                if "blobSize" in item_dict:
                    item_dict["formatted_size"] = format_file_size(
                        item_dict.get("blobSize")
                    )

                # Get status from cached transcript statuses
                if "transcriptId" in item_dict:
                    item_dict["status"] = transcript_statuses.get(
                        item_dict.get("transcriptId"), "error"
                    )
                    logger.debug(
                        "Transcript status retrieved",
                        extra={
                            "transcript_id": item_dict.get("transcriptId"),
                            "status": item_dict.get("status"),
                        },
                    )
                else:
                    item_dict["status"] = "pending"
                    logger.debug(
                        "No transcript ID found for item",
                        extra={"row_key": item_dict.get("RowKey", "unknown")},
                    )

                item_dict["_previous_status"] = item_dict.get("status")

                # Process timestamp
                if "uploadTime" not in item_dict:
                    item_dict["uploadTime"] = item_dict.get("Timestamp", MIN_DATE)
                    logger.debug(
                        "Using fallback timestamp",
                        extra={"row_key": item_dict.get("RowKey", "unknown")},
                    )

                try:
                    # Handle different timestamp types
                    if isinstance(item_dict["uploadTime"], str):
                        dt = datetime.fromisoformat(
                            item_dict["uploadTime"].replace("Z", "+00:00")
                        )
                    elif isinstance(item_dict["uploadTime"], datetime):
                        dt = item_dict["uploadTime"]
                        if dt.tzinfo is None:
                            dt = dt.replace(tzinfo=pytz.UTC)
                    else:
                        dt = MIN_DATE

                    local_dt = dt.astimezone(
                        pytz.timezone(st.session_state["timezone"])
                    )
                    item_dict["_timestamp"] = local_dt
                    item_dict["uploadTime"] = local_dt
                except ValueError as e:
                    logger.error(
                        "Error parsing timestamp",
                        exc_info=True,
                        extra={
                            "row_key": item_dict.get("RowKey", "unknown"),
                            "upload_time": item_dict.get("uploadTime"),
                            "error": str(e),
                        },
                    )
                    item_dict["_timestamp"] = MIN_DATE
                    item_dict["uploadTime"] = MIN_DATE

                items_list.append(item_dict)

            except Exception as e:
                logger.error(
                    "Error processing table item",
                    exc_info=True,
                    extra={"error": str(e), "row_key": item.get("RowKey", "unknown")},
                )
                continue

        logger.info(
            "Table data load completed",
            extra={"processed_items": len(items_list), "total_items": len(items)},
        )
        return items_list

    def display_transcript_item(item):
        """Display a single transcript item in a fragment"""
        try:
            transcript_id = item.get("transcriptId")
            if not transcript_id:
                logger.error(
                    "Missing transcript ID",
                    extra={"row_key": item.get("RowKey", "unknown")},
                )
                return

            logger.debug(
                "Starting to display transcript item",
                extra={
                    "transcript_id": transcript_id,
                    "status": item.get("status"),
                    "row_key": item.get("RowKey", "unknown"),
                },
            )

            # Get status info for formatting
            upload_time = item.get("uploadTime")
            upload_time_str = localized_timestamp(upload_time)
            original_file_name = item.get("originalFileName", "Untitled")
            row_key = item.get("RowKey", "")
            status = item.get("status")
            class_name = item.get("className", "")
            uploader_email = item.get("uploaderEmail", "Unknown")

            logger.debug(
                "Retrieved item metadata",
                extra={
                    "upload_time": upload_time_str,
                    "file_name": original_file_name,
                    "status": status,
                },
            )

            # Choose icon based on status
            status_icon = "📄"  # Default icon
            if status in ["queued", "processing"]:
                status_icon = "⏳"
            elif status in ["error", "failed"]:
                status_icon = "❌"

            # Format display name - use class name if available, otherwise use original file name
            display_name = class_name if class_name else original_file_name

            # Format uploader email - show only the part before @ symbol
            uploader_display = (
                uploader_email.split("@")[0]
                if "@" in uploader_email
                else uploader_email
            )

            # Format date to be more concise
            date_display = (
                upload_time.strftime("%Y-%m-%d")
                if isinstance(upload_time, datetime)
                else str(upload_time)
            )

            logger.debug("Starting to render expander UI component")
            with st.expander(
                f"{status_icon} {display_name} | by {uploader_display} | {date_display}",
                expanded=False,
            ):
                description = item.get("description", "")
                size = item.get("formatted_size", "")

                if class_name:
                    st.write(f"**Class**: {class_name}")
                    st.write(f"**File**: {original_file_name}")
                else:
                    st.write(f"**File**: {original_file_name}")

                if description:
                    st.write(f"**Description**: {description}")
                st.write(f"**Uploaded**: {upload_time_str}")
                st.write(f"**Size**: {size}")

                logger.debug("Starting to process audio URL")
                # Audio player
                audio_url_with_sas = get_sas_url_for_audio_file_name(row_key)
                if audio_url_with_sas:
                    st.audio(audio_url_with_sas)

                # Only fetch full transcript details if status is completed and user expands the item
                transcript = None
                if status == "completed" and transcript_id:
                    logger.debug(
                        f"Fetching completed transcript for ID: {transcript_id}"
                    )
                    try:
                        transcript = aai.Transcript.get_by_id(transcript_id)
                        logger.debug(
                            "Successfully retrieved transcript from AssemblyAI"
                        )
                    except Exception as e:
                        if "api error" in str(e).lower():
                            logger.error(
                                "AssemblyAI API error",
                                extra={"transcript_id": transcript_id, "error": str(e)},
                                exc_info=True,
                            )
                            st.info("Transcript temporarily unavailable")
                        elif isinstance(e, pydantic.ValidationError):
                            logger.warning(
                                "Validation error for transcript",
                                extra={"transcript_id": transcript_id, "error": str(e)},
                            )
                            try:
                                transcript = aai.Transcript.get_by_id(transcript_id)
                                if (
                                    not hasattr(transcript, "text")
                                    or not transcript.text
                                ):
                                    transcript = None
                                    logger.warning(
                                        "Secondary transcript fetch returned invalid data"
                                    )
                            except Exception as inner_e:
                                logger.error(
                                    "Secondary transcript fetch error",
                                    extra={
                                        "transcript_id": transcript_id,
                                        "error": str(inner_e),
                                    },
                                    exc_info=True,
                                )
                                transcript = None
                        else:
                            logger.error(
                                "Unexpected error loading transcript",
                                extra={"transcript_id": transcript_id, "error": str(e)},
                                exc_info=True,
                            )
                            st.info("Unable to load transcript")

                    if transcript:
                        logger.debug("Starting to render transcript download buttons")
                        # Add download buttons in a row
                        col1, col2 = st.columns([1, 1])

                        with col1:
                            # Create download button for markdown
                            full_markdown = generate_transcript_markdown(transcript)
                            st.download_button(
                                label="Download as Markdown",
                                data=full_markdown,
                                file_name=f"{original_file_name}.md",
                                mime="text/markdown",
                                key=f"download_transcript_md_{row_key}",
                            )

                        with col2:
                            # Create download button for docx
                            docx_bytes = generate_transcript_docx(transcript)
                            st.download_button(
                                label="Download as Word",
                                data=docx_bytes,
                                file_name=f"{original_file_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_transcript_docx_{row_key}",
                            )

                        logger.debug("Starting to render transcript content")
                        ### Show transcript
                        st.markdown("#### 📝 Transcript")
                        full_markdown = generate_transcript_markdown(transcript)
                        st.markdown(full_markdown)
                        logger.debug("Completed rendering transcript content")

                elif status in ["queued", "processing"]:
                    with st.container(border=True):
                        st.markdown(
                            """
                            #### ⏳ Processing
                            The transcript is still being generated. This typically takes 1-2 minutes.
                            """
                        )
                        if st.button(
                            "Refresh Status",
                            icon="🔄",
                            key=f"refresh_status_body_{row_key}",
                        ):
                            st.rerun()

                elif status in ["error", "failed"]:
                    with st.container(border=True):
                        st.markdown(
                            """
                            #### ❌ Error
                            There was a problem processing this transcript. Please try uploading the file again.
                            """
                        )

            logger.debug(
                "Completed displaying transcript item",
                extra={"transcript_id": transcript_id, "row_key": row_key},
            )

        except pydantic.ValidationError as e:
            logger.error(
                "Validation error displaying transcript",
                extra={
                    "transcript_id": transcript_id,
                    "error": str(e),
                    "row_key": item.get("RowKey", "unknown"),
                },
                exc_info=True,
            )
            st.error(
                f"Could not load transcript {transcript_id} - the data may be corrupted or deleted"
            )

        except Exception as e:
            logger.error(
                "Error displaying transcript",
                extra={
                    "transcript_id": transcript_id,
                    "error": str(e),
                    "row_key": item.get("RowKey", "unknown"),
                },
                exc_info=True,
            )
            st.error(f"Error loading transcript: {str(e)}")

    def display_table_data():
        """Display the table data with progress indicators and error handling"""
        try:
            logger.info("Starting to load and display table data")
            items_list = load_table_data(table_client)

            if not items_list:
                logger.info("No files found in table data")
                st.info("No files found")
                return

            # Sort by timestamp (newest first)
            items_list.sort(
                key=lambda x: x.get("_timestamp", datetime.min), reverse=True
            )
            logger.debug(
                "Sorted table data",
                extra={
                    "total_items": len(items_list),
                    "first_item_time": (
                        items_list[0].get("_timestamp") if items_list else None
                    ),
                },
            )

            # Calculate pagination
            total_items = len(items_list)
            start_idx = 0
            end_idx = st.session_state["items_per_page"]

            logger.debug(
                "Pagination calculated",
                extra={
                    "total_items": total_items,
                    "items_per_page": st.session_state["items_per_page"],
                    "start_idx": start_idx,
                    "end_idx": end_idx,
                },
            )

            # Display items in fragments
            for idx, item in enumerate(items_list[start_idx:end_idx]):
                try:
                    logger.debug(
                        f"Rendering item {idx + 1} of {min(end_idx, total_items)}",
                        extra={"row_key": item.get("RowKey", "unknown")},
                    )
                    with st.container():
                        display_transcript_item(item)
                except Exception as e:
                    logger.error(
                        "Error displaying item",
                        extra={
                            "error": str(e),
                            "row_key": item.get("RowKey", "unknown"),
                            "index": idx,
                        },
                        exc_info=True,
                    )
                    st.error(f"Error displaying item: {str(e)}")
                    continue

            # Load More button
            if end_idx < total_items:
                logger.debug("Rendering Load More button")
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button(
                        f"Load More ({total_items - end_idx} remaining)",
                        use_container_width=True,
                    ):
                        logger.info("Load More button clicked")
                        st.session_state["items_per_page"] += 5
                        st.rerun()

            # Show total count
            st.caption(
                f"Showing {min(end_idx, total_items)} of {total_items} transcripts"
            )
            logger.info(
                "Completed displaying table data",
                extra={
                    "displayed_items": min(end_idx, total_items),
                    "total_items": total_items,
                },
            )

        except Exception as e:
            logger.error(
                "Error in display_table_data", extra={"error": str(e)}, exc_info=True
            )
            st.error("Unable to load transcripts. Please try again later.")

    display_table_data()

except Exception as e:
    logger.error(f"Fatal initialization error: {str(e)}", exc_info=True)
    st.error("Unable to start application. Please try again later.")
    st.stop()
